import docx
import os

def analyze_docx_format(docx_path):
    """分析Word文档的排版格式"""
    doc = docx.Document(docx_path)
    
    # 提取文档信息
    format_info = {
        "filename": os.path.basename(docx_path),
        "page_setup": {},
        "styles": {},
        "paragraphs": [],
        "tables": []
    }
    
    # 页面设置（注意：python-docx对页面设置的支持有限）
    try:
        section = doc.sections[0]
        format_info["page_setup"]["page_width"] = section.page_width.inches
        format_info["page_setup"]["page_height"] = section.page_height.inches
        format_info["page_setup"]["left_margin"] = section.left_margin.inches
        format_info["page_setup"]["right_margin"] = section.right_margin.inches
        format_info["page_setup"]["top_margin"] = section.top_margin.inches
        format_info["page_setup"]["bottom_margin"] = section.bottom_margin.inches
    except Exception as e:
        print(f"提取页面设置时出错: {e}")
    
    # 提取样式信息
    for style in doc.styles:
        if style.type == docx.enum.style.WD_STYLE_TYPE.PARAGRAPH:
            format_info["styles"][style.name] = {
                "font_name": style.font.name if style.font.name else "",
                "font_size": style.font.size.pt if style.font.size else "",
                "bold": style.font.bold,
                "italic": style.font.italic,
                "color": style.font.color.rgb if style.font.color.rgb else "",
                "alignment": style.paragraph_format.alignment
            }
    
    # 提取段落格式
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            para_info = {
                "text": para.text[:100] + "..." if len(para.text) > 100 else para.text,
                "style": para.style.name,
                "font_name": para.runs[0].font.name if para.runs else "",
                "font_size": para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else "",
                "bold": para.runs[0].font.bold if para.runs else False,
                "italic": para.runs[0].font.italic if para.runs else False,
                "alignment": para.alignment,
                "line_spacing": para.paragraph_format.line_spacing
            }
            format_info["paragraphs"].append(para_info)
    
    # 提取表格信息
    for i, table in enumerate(doc.tables):
        table_info = {
            "rows": len(table.rows),
            "columns": len(table.columns),
            "cells": []
        }
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cell_info = {
                        "text": cell.text[:50] + "..." if len(cell.text) > 50 else cell.text,
                        "font_name": cell.paragraphs[0].runs[0].font.name if cell.paragraphs and cell.paragraphs[0].runs else "",
                        "font_size": cell.paragraphs[0].runs[0].font.size.pt if cell.paragraphs and cell.paragraphs[0].runs and cell.paragraphs[0].runs[0].font.size else ""
                    }
                    table_info["cells"].append(cell_info)
        format_info["tables"].append(table_info)
    
    return format_info

def main():
    docx_files = [
        r"y:\Trae_Solo\documents\“双碳”新能源领域月度行业动态 2026年1月(1).docx",
        r"y:\Trae_Solo\documents\“双碳”新能源领域月度行业动态 2026年2月(1).docx",
        r"y:\Trae_Solo\documents\“双碳”新能源领域月度行业动态 2026年3月(1).docx"
    ]
    
    all_format_info = []
    for docx_file in docx_files:
        print(f"分析文件: {docx_file}")
        format_info = analyze_docx_format(docx_file)
        all_format_info.append(format_info)
    
    # 生成skill文件内容
    skill_content = generate_skill_file(all_format_info)
    
    # 保存skill文件
    skill_file_path = r"y:\Trae_Solo\word_format_skill.json"
    with open(skill_file_path, 'w', encoding='utf-8') as f:
        import json
        json.dump(skill_content, f, ensure_ascii=False, indent=2)
    
    print(f"\nSkill文件已生成: {skill_file_path}")

def generate_skill_file(format_infos):
    """根据分析结果生成skill文件"""
    # 提取共同格式
    common_styles = {}
    if format_infos:
        # 以第一个文档为基础
        base_styles = format_infos[0]["styles"]
        # 检查其他文档是否有相同的样式
        for style_name, style_info in base_styles.items():
            common = True
            for info in format_infos[1:]:
                if style_name not in info["styles"] or info["styles"][style_name] != style_info:
                    common = False
                    break
            if common:
                common_styles[style_name] = style_info
    
    # 生成skill文件结构
    skill = {
        "name": "word_format_analyzer",
        "description": "分析Word文档排版格式并应用到新文档",
        "version": "1.0.0",
        "author": "Trae AI",
        "date": "2026-04-16",
        "formats": {
            "page_setup": format_infos[0]["page_setup"] if format_infos else {},
            "common_styles": common_styles,
            "sample_paragraphs": format_infos[0]["paragraphs"][:5] if format_infos and format_infos[0]["paragraphs"] else [],
            "tables": format_infos[0]["tables"] if format_infos else []
        },
        "instructions": [
            "打开Word文档",
            "应用页面设置",
            "创建并应用样式",
            "设置段落格式",
            "应用表格格式"
        ]
    }
    
    return skill

if __name__ == "__main__":
    main()